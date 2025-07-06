import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def read_excel_data(excel_file, sheet_name=None):
    """
    从Excel文件读取数据
    
    Args:
        excel_file (str): Excel文件路径
        sheet_name (str): 工作表名称，如果为None则读取第一个工作表
    
    Returns:
        pandas.DataFrame: 读取的数据
    """
    try:
        if sheet_name:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
        else:
            df = pd.read_excel(excel_file)
        return df
    except Exception as e:
        print(f"读取Excel文件时出错: {e}")
        return None

def create_word_document(data, output_file, template_type="report"):
    """
    根据Excel数据创建Word文档
    
    Args:
        data (pandas.DataFrame): Excel数据
        output_file (str): 输出Word文件路径
        template_type (str): 模板类型 ("report", "table", "list")
    """
    try:
        # 创建Word文档
        doc = Document()
        
        if template_type == "report":
            create_report_format(doc, data)
        elif template_type == "table":
            create_table_format(doc, data)
        elif template_type == "list":
            create_list_format(doc, data)
        else:
            create_report_format(doc, data)  # 默认使用报告格式
        
        # 保存文档
        doc.save(output_file)
        print(f"Word文档已生成: {output_file}")
        
    except Exception as e:
        print(f"生成Word文档时出错: {e}")

def create_report_format(doc, data):
    """创建报告格式的Word文档"""
    # 添加标题
    title = doc.add_heading('数据报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加日期
    from datetime import datetime
    date_para = doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y年%m月%d日")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 添加概述
    doc.add_heading('数据概述', level=1)
    overview = doc.add_paragraph()
    overview.add_run(f'本报告包含 {len(data)} 行数据，共 {len(data.columns)} 个字段。')
    
    # 按行生成内容
    doc.add_heading('详细数据', level=1)
    
    for index, row in data.iterrows():
        # 为每行数据创建一个小节
        doc.add_heading(f'记录 {index + 1}', level=2)
        
        # 添加该行的所有字段信息
        for col in data.columns:
            para = doc.add_paragraph()
            para.add_run(f'{col}: ').bold = True
            para.add_run(str(row[col]))

def create_table_format(doc, data):
    """创建表格格式的Word文档"""
    # 添加标题
    title = doc.add_heading('数据表格', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 创建表格
    table = doc.add_table(rows=1, cols=len(data.columns))
    table.style = 'Table Grid'
    
    # 设置表头
    header_cells = table.rows[0].cells
    for i, column in enumerate(data.columns):
        header_cells[i].text = str(column)
        # 设置表头加粗
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # 添加数据行
    for index, row in data.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

def create_list_format(doc, data):
    """创建列表格式的Word文档"""
    # 添加标题
    title = doc.add_heading('数据列表', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 按行创建编号列表
    for index, row in data.iterrows():
        # 添加主项目
        main_para = doc.add_paragraph(f'项目 {index + 1}', style='List Number')
        
        # 添加子项目（各字段）
        for col in data.columns:
            sub_para = doc.add_paragraph(f'{col}: {row[col]}', style='List Bullet')

def batch_process_excel_files(input_folder, output_folder, template_type="report"):
    """
    批量处理文件夹中的Excel文件
    
    Args:
        input_folder (str): 输入文件夹路径
        output_folder (str): 输出文件夹路径
        template_type (str): 模板类型
    """
    # 确保输出文件夹存在
    os.makedirs(output_folder, exist_ok=True)
    
    # 获取所有Excel文件
    excel_files = [f for f in os.listdir(input_folder) 
                   if f.endswith(('.xlsx', '.xls'))]
    
    for excel_file in excel_files:
        input_path = os.path.join(input_folder, excel_file)
        output_filename = os.path.splitext(excel_file)[0] + '.docx'
        output_path = os.path.join(output_folder, output_filename)
        
        print(f"正在处理: {excel_file}")
        
        # 读取Excel数据
        data = read_excel_data(input_path)
        if data is not None:
            # 生成Word文档
            create_word_document(data, output_path, template_type)

# 使用示例
def main():
    """主函数 - 使用示例"""
    
    # 单个文件处理示例
    excel_file = "data.xlsx"  # 替换为你的Excel文件路径
    output_file = "output_report.docx"  # 输出Word文件路径
    
    # 检查文件是否存在
    if os.path.exists(excel_file):
        # 读取Excel数据
        data = read_excel_data(excel_file)
        
        if data is not None:
            print("Excel数据预览:")
            print(data.head())
            print(f"\n数据形状: {data.shape}")
            
            # 生成不同格式的Word文档
            create_word_document(data, "report_format.docx", "report")
            create_word_document(data, "table_format.docx", "table")  
            create_word_document(data, "list_format.docx", "list")
    else:
        print(f"Excel文件不存在: {excel_file}")
        print("请确保文件路径正确，或创建示例数据...")
        
        # 创建示例数据
        create_sample_data()

def create_sample_data():
    """创建示例Excel数据用于测试"""
    sample_data = {
        '姓名': ['张三', '李四', '王五', '赵六'],
        '年龄': [25, 30, 28, 35],
        '部门': ['技术部', '销售部', '人事部', '财务部'],
        '工资': [8000, 9000, 7500, 8500],
        '入职日期': ['2023-01-15', '2022-06-10', '2023-03-20', '2021-12-05']
    }
    
    df = pd.DataFrame(sample_data)
    df.to_excel('sample_data.xlsx', index=False)
    print("已创建示例Excel文件: sample_data.xlsx")
    
    # 使用示例数据生成Word文档
    create_word_document(df, "sample_report.docx", "report")
    create_word_document(df, "sample_table.docx", "table")
    create_word_document(df, "sample_list.docx", "list")

if __name__ == "__main__":
    main()